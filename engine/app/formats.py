"""Multi-format intake: PDF, images, and Word documents.

Everything is normalised to a **PDF rendition** at intake. Downstream — page
rendering, the bbox overlay, the document viewer, vision extraction — then has
exactly one shape to handle, and adding a format later means adding one
converter here rather than a branch in five modules.

The original bytes are always kept alongside the rendition. The rendition is
what a reviewer looks at and what the model reads; the original is what the
SHA-256 idempotency key is computed over, so re-submitting the same file is
still caught however it was produced.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Optional, Tuple

import fitz  # PyMuPDF

log = logging.getLogger("engine.formats")

# A4 at 72dpi, for laying out documents that have no page geometry of their own.
PAGE_W, PAGE_H = 595.0, 842.0
MARGIN = 48.0


@dataclass
class Format:
    key: str
    label: str
    mime: str
    extensions: tuple
    needs_ocr: bool          # True when the format carries no text layer


PDF = Format("PDF", "PDF document", "application/pdf", (".pdf",), False)
IMAGE = Format("IMAGE", "Image", "image/*",
               (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"),
               True)
DOCX = Format("DOCX", "Word document",
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              (".docx",), False)
DOC_LEGACY = Format("DOC", "Word 97-2003 document", "application/msword", (".doc",), False)
UNKNOWN = Format("UNKNOWN", "Unrecognised file", "application/octet-stream", (), False)

ACCEPTED_EXTENSIONS = PDF.extensions + IMAGE.extensions + DOCX.extensions
ACCEPTED_MIME = (
    "application/pdf,image/png,image/jpeg,image/tiff,image/bmp,image/webp,"
    "image/gif,application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def sniff(data: bytes, filename: str = "") -> Format:
    """Identify the format from magic bytes, falling back to the extension.

    Magic bytes first on purpose: a file's real type matters more than what it
    was named, and a mislabelled upload should be handled correctly rather than
    rejected on a technicality.
    """
    if data[:5] == b"%PDF-":
        return PDF
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return IMAGE
    if data[:3] == b"\xff\xd8\xff":                      # JPEG
        return IMAGE
    if data[:4] in (b"II*\x00", b"MM\x00*"):             # TIFF
        return IMAGE
    if data[:2] == b"BM":                                # BMP
        return IMAGE
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return IMAGE
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return IMAGE
    if data[:4] == b"PK\x03\x04":
        # A zip container — .docx and .xlsx both look like this. Peek inside.
        try:
            import zipfile

            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = set(zf.namelist())
            if "word/document.xml" in names:
                return DOCX
        except Exception:
            pass
        return UNKNOWN
    if data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":  # OLE2 — legacy .doc
        return DOC_LEGACY

    lower = (filename or "").lower()
    for fmt in (PDF, IMAGE, DOCX):
        if any(lower.endswith(ext) for ext in fmt.extensions):
            return fmt
    return UNKNOWN


# ----------------------------------------------------------------------
def to_pdf(data: bytes, filename: str = "") -> Tuple[bytes, Format, Optional[str]]:
    """Return ``(pdf_bytes, detected_format, error)``.

    ``error`` is a human-readable reason when conversion is impossible; the
    caller turns that into an ING-01 failure rather than a crash.
    """
    fmt = sniff(data, filename)

    if fmt is PDF:
        return data, fmt, None
    if fmt is IMAGE:
        try:
            return _image_to_pdf(data), fmt, None
        except Exception as exc:
            log.warning("Image conversion failed for %s: %s", filename, exc)
            return b"", fmt, f"The image could not be read ({exc})."
    if fmt is DOCX:
        try:
            return _docx_to_pdf(data), fmt, None
        except ImportError:
            return b"", fmt, ("Word support needs python-docx. Install it with "
                              "`pip install python-docx`.")
        except Exception as exc:
            log.warning("DOCX conversion failed for %s: %s", filename, exc)
            return b"", fmt, f"The Word document could not be read ({exc})."
    if fmt is DOC_LEGACY:
        return b"", fmt, (
            "Legacy Word 97-2003 (.doc) is not supported. Save it as .docx or "
            "print it to PDF and upload that."
        )

    return b"", fmt, (
        "Unrecognised file type. Upload a PDF, an image (PNG, JPEG, TIFF, BMP, "
        "WEBP), or a Word .docx document."
    )


def _image_to_pdf(data: bytes) -> bytes:
    """Wrap an image in a single PDF page, sized to the image's aspect ratio.

    Scanned invoices arrive as photos and screenshots constantly. Preserving the
    aspect ratio matters because the bbox overlay is expressed as a fraction of
    the page — stretching the image would put every highlight in the wrong place.
    """
    from PIL import Image, ImageOps

    with Image.open(io.BytesIO(data)) as img:
        # Phone cameras record rotation in EXIF rather than in the pixels.
        img = ImageOps.exif_transpose(img)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        width, height = img.size
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        png = buffer.getvalue()

    scale = min(PAGE_W / width, PAGE_H / height)
    draw_w, draw_h = width * scale, height * scale
    x, y = (PAGE_W - draw_w) / 2, (PAGE_H - draw_h) / 2

    doc = fitz.open()
    page = doc.new_page(width=PAGE_W, height=PAGE_H)
    page.insert_image(fitz.Rect(x, y, x + draw_w, y + draw_h), stream=png)
    try:
        doc.set_metadata({"producer": "invoice-platform", "creator": "",
                          "title": "", "author": "", "subject": "",
                          "keywords": "", "creationDate": "", "modDate": ""})
        return doc.tobytes(garbage=4, deflate=True)
    finally:
        doc.close()


def _docx_to_pdf(data: bytes) -> bytes:
    """Lay a Word document out as a PDF, preserving paragraphs and tables.

    Not a faithful renderer — it does not attempt fonts, images or styling. It
    exists so the text is *positioned* on a page, which is what gives the
    extraction real word geometry and lets the bbox overlay work on a .docx the
    same way it does on a PDF.
    """
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    blocks = _docx_blocks(document)

    out = fitz.open()
    page = out.new_page(width=PAGE_W, height=PAGE_H)
    y = MARGIN

    def new_page():
        nonlocal page, y
        page = out.new_page(width=PAGE_W, height=PAGE_H)
        y = MARGIN

    for block in blocks:
        if block["kind"] == "table":
            for row in block["rows"]:
                if y > PAGE_H - MARGIN:
                    new_page()
                x = MARGIN
                cell_w = (PAGE_W - 2 * MARGIN) / max(len(row), 1)
                for cell in row:
                    page.insert_text((x, y), _clip(cell, cell_w),
                                     fontname="helv", fontsize=8)
                    x += cell_w
                y += 13
            y += 6
            continue

        text = block["text"]
        if not text:
            y += 7
            continue

        size = block["size"]
        font = "hebo" if block["bold"] else "helv"
        for line in _wrap(text, PAGE_W - 2 * MARGIN, size):
            if y > PAGE_H - MARGIN:
                new_page()
            page.insert_text((MARGIN, y), line, fontname=font, fontsize=size)
            y += size + 4
        y += 3

    try:
        out.set_metadata({"producer": "invoice-platform", "creator": "",
                          "title": "", "author": "", "subject": "",
                          "keywords": "", "creationDate": "", "modDate": ""})
        return out.tobytes(garbage=4, deflate=True)
    finally:
        out.close()


def _docx_blocks(document) -> list:
    """Walk the document body in order, so tables stay where the author put
    them. ``document.paragraphs`` and ``document.tables`` are separate lists and
    iterating them one after the other scrambles an invoice's layout."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    blocks = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para = Paragraph(child, document)
            style = (para.style.name or "").lower()
            heading = "heading" in style or "title" in style
            bold = heading or any(r.bold for r in para.runs if r.bold is not None)
            blocks.append({
                "kind": "paragraph",
                "text": para.text.strip(),
                "size": 13 if "title" in style else 11 if heading else 9,
                "bold": bool(bold),
            })
        elif tag == "tbl":
            table = Table(child, document)
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            blocks.append({"kind": "table", "rows": rows})
    return blocks


def _wrap(text: str, width: float, size: float) -> list:
    """Greedy wrap using Helvetica metrics so the layout matches what PyMuPDF
    will actually draw."""
    words = text.split()
    if not words:
        return []
    lines, current = [], words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if fitz.get_text_length(candidate, fontname="helv", fontsize=size) <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _clip(text: str, width: float, size: float = 8) -> str:
    text = text.replace("\n", " ")
    if fitz.get_text_length(text, fontname="helv", fontsize=size) <= width - 6:
        return text
    while text and fitz.get_text_length(text + "…", fontname="helv",
                                        fontsize=size) > width - 6:
        text = text[:-1]
    return text + "…"
